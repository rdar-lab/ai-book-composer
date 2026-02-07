"""Enhanced tools with security, logging, and additional format support."""

import logging
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..config import Settings

logger = logging.getLogger(__name__)


class BookWriter:
    """Tool to generate final book in DOCX format."""

    def __init__(self, settings: Settings, output_directory: str):
        self.settings = settings
        self.output_directory = Path(output_directory).resolve()
        self.output_directory.mkdir(parents=True, exist_ok=True)
        logger.info(f"BookGeneratorTool initialized for directory: {self.output_directory}")

    def run(
            self,
            title: str,
            author: str,
            chapters: List[Dict[str, Any]],
            references: List[str],
            output_filename: str = "book.docx"
    ) -> Dict[str, Any]:
        """Generate final book in DOCX format.

        Args:
            title: Book title
            author: Book author
            chapters: List of chapter dictionaries with 'title', 'content', and optionally 'images'
            references: List of reference strings
            output_filename: Output filename

        Returns:
            Dictionary with file path and status
        """
        logger.info(f"Generating book: {title} by {author}")
        logger.debug(f"Chapters: {len(chapters)}, References: {len(references)}")

        try:
            file_path = self.output_directory / output_filename

            # Create document
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Title page
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(24)
            title_run.font.name = 'Arial'
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Add spacing
            
            author_para = doc.add_paragraph()
            author_run = author_para.add_run(f"By {author}")
            author_run.font.size = Pt(14)
            author_run.font.name = 'Arial'
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(datetime.now().strftime("%B %Y"))
            date_run.font.size = Pt(12)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Table of Contents (on new page)
            doc.add_page_break()
            
            toc_heading = doc.add_heading("Table of Contents", level=1)
            
            for i, chapter in enumerate(chapters, 1):
                chapter_title = chapter.get("title", f"Chapter {i}")
                toc_entry = doc.add_paragraph(f"Chapter {i}: {chapter_title}")

            # Chapters
            total_images_added = 0
            for i, chapter in enumerate(chapters, 1):
                chapter_title = chapter.get("title", f"Chapter {i}")
                chapter_content = chapter.get("content", "")
                chapter_images = chapter.get("images", [])

                # Chapter heading (on new page)
                doc.add_page_break()
                doc.add_heading(f"Chapter {i}: {chapter_title}", level=1)

                # Categorize images by position
                start_images = []
                middle_images = []
                end_images = []

                for img in chapter_images:
                    position = str(img.get("position", "middle")).lower()
                    if position == "start":
                        start_images.append(img)
                    elif position == "end":
                        end_images.append(img)
                    else:
                        middle_images.append(img)

                # Add start images
                for img in start_images:
                    if self._add_image_to_doc(doc, img):
                        total_images_added += 1

                # Chapter content - split into paragraphs
                paragraphs = chapter_content.split('\n\n')
                num_paragraphs = len([p for p in paragraphs if p.strip()])

                # Determine where to insert middle images
                middle_insert_point = num_paragraphs // 2 if num_paragraphs > 0 else 0

                for para_idx, para_text in enumerate(paragraphs):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())

                        # Insert middle images at approximately the middle of the chapter
                        if para_idx == middle_insert_point and middle_images:
                            for img in middle_images:
                                if self._add_image_to_doc(doc, img):
                                    total_images_added += 1

                # Add end images
                for img in end_images:
                    if self._add_image_to_doc(doc, img):
                        total_images_added += 1

            # References
            if references:
                doc.add_page_break()
                doc.add_heading("References", level=1)

                for ref in references:
                    doc.add_paragraph(ref)

            # Save document
            doc.save(str(file_path))

            logger.info(f"Book generated successfully: {file_path} with {total_images_added} images")

            return {
                "success": True,
                "file_path": str(file_path),
                "chapter_count": len(chapters),
                "reference_count": len(references),
                "image_count": total_images_added
            }
        except Exception as e:
            logger.exception(f"Error generating book: {e}")
            return {
                "success": False,
                "error": str(e)
            }

    @staticmethod
    def _add_image_to_doc(doc: Document, img_info: Dict[str, Any]) -> bool:
        """Add an image to the document.

        Args:
            doc: DOCX document to add image to
            img_info: Dictionary with image information (path, reasoning, etc.)
            
        Returns:
            True if image was successfully added, False otherwise
        """
        image_path = img_info.get("image_path", "")

        if not image_path or not Path(image_path).exists():
            logger.warning(f"Image not found: {image_path}")
            return False

        try:
            # Add image with reasonable width (6 inches, maintains aspect ratio)
            doc.add_picture(str(image_path), width=Inches(6))
            
            # Add caption if reasoning is provided
            reasoning = img_info.get("reasoning", "")
            if reasoning:
                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(f"Figure: {reasoning}")
                caption_run.italic = True
                caption_run.font.size = Pt(10)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            logger.debug(f"Added image to document: {image_path}")
            return True
        except Exception as img_error:
            logger.exception(f"Could not add image {image_path}: {repr(img_error)}")
            return False
